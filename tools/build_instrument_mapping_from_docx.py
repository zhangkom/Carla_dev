# /**
# * File name: build_instrument_mapping_from_docx.py
# * Brief: 云端 DAW 音源映射配置生成工具
# * Function:
# *     从云端 DAW 需求文档中提取 MIDI 音源映射并生成结构化 JSON 配置
# * Author: 咪咕数创工程架构组
# *     MGSC AI Software Architecture group
# * Version: V2.5.10
# * Date: 2026/04/30
# */
import argparse
import json
import sys
from collections import Counter
from datetime import date
from pathlib import Path


DOC_NAME = "云端DAW音频工作站引擎.docx"
OUTPUT_NAME = "instrument_mapping.deploy.json"


PLUGIN_INFO = {
    "Musyng_Kite": {
        "plugin_id": "sf2_musyng_kite",
        "plugin_type": "sf2",
        "asset_relpath": "soundfont2/Musyng_Kite.sf2",
        "implementation_status": "implemented",
    },
    "Keyzone Classic": {
        "plugin_id": "vst_keyzone_classic",
        "plugin_type": "vst2",
        "asset_relpath": "Steinberg/VstPlugins/Keyzone Classic/Keyzone Classic.dll",
        "implementation_status": "planned",
    },
    "Sonatina Orchestra": {
        "plugin_id": "vst_sonatina_orchestra",
        "plugin_type": "vst2",
        "asset_relpath": "Steinberg/VstPlugins/Sonatina Orchestra/Sonatina Orchestra.dll",
        "implementation_status": "planned",
    },
    "DSK Saxophones": {
        "plugin_id": "vst_dsk_saxophones",
        "plugin_type": "vst2",
        "asset_relpath": "Steinberg/VstPlugins/DSK Saxophones/DSK Saxophones.dll",
        "implementation_status": "planned",
    },
    "kong": {
        "plugin_id": "kong_qin_rv",
        "plugin_type": "vst2",
        "asset_relpath": "kong_audio/qin_rv_v2_2/vst2/Qin_RV.DLL",
        "implementation_status": "partial",
    },
}


COMPATIBILITY_OVERRIDES = {
    ("15", "kong"): {
        "normalized_bank": "ChineeYangQin",
        "local_preset_relpath": "kong_audio/qin_rv_v2_2/library/ChineeYangQin",
        "needs_confirmation": [
            "文档 bank=yangqin 是逻辑名，不是本地真实目录名",
            "当前服务尚未建立 Kong YangQin 状态文件",
        ],
    },
    ("40", "Sonatina Orchestra"): {
        "normalized_bank": "Sonatina Violin",
        "notes": ["文档写为 Sonatina violin，本地目录为 Sonatina Violin"],
    },
    ("46", "Sonatina Orchestra"): {
        "normalized_program": "Default Group",
        "notes": ["文档写为 default group，本地文件为 Default Group.txt"],
    },
    ("57", "Sonatina Orchestra"): {
        "normalized_program": "Tenor Trombone",
        "needs_confirmation": ["文档写为 Tensor Trombone，疑似应为 Tenor Trombone"],
    },
    ("58", "Sonatina Orchestra"): {
        "normalized_program": "Tuba Sustain",
        "needs_confirmation": ["文档写为 Tuba  / Sustain，本地文件为 Tuba Sustain.txt"],
    },
    ("60", "Sonatina Orchestra"): {
        "normalized_program": "Solo Horn",
        "needs_confirmation": ["文档写为 Solo  / Horn，本地文件为 Solo Horn.txt"],
    },
    ("64", "DSK Saxophones"): {
        "normalized_program": "Soprano Sax",
        "needs_confirmation": ["文档写为 Soprano /  Sax，本地文件为 Soprano Sax.txt"],
    },
    ("65", "Musyng_Kite"): {
        "needs_confirmation": [
            "文档中 MIDI id=65，云端 program=65，但 web program 写为 63",
        ],
    },
    ("70", "Sonatina Orchestra"): {
        "normalized_bank": "Sonatina Bassoon",
        "needs_confirmation": ["文档写为 Sonatina Bassoom，疑似应为 Sonatina Bassoon"],
    },
    ("71", "Sonatina Orchestra"): {
        "normalized_program": "Solo Clarinet",
        "needs_confirmation": ["文档写为 Solo / Clarinet，本地文件为 Solo Clarinet.txt"],
    },
    ("107", "kong"): {
        "normalized_bank": "ChineeGuZheng_Classic",
        "local_preset_relpath": "kong_audio/qin_rv_v2_2/library/ChineeGuZheng_Classic",
        "candidate_relpaths": [
            "kong_audio/qin_rv_v2_2/library/ChineeGuZheng_Classic",
            "kong_audio/qin_rv_v2_2/library/ChineeGuZheng_II",
        ],
        "needs_confirmation": [
            "文档 bank=chineseGuZheng 是逻辑名，不是本地真实目录名",
            "需要确认最终使用 ChineeGuZheng_Classic 还是 ChineeGuZheng_II",
            "当前服务尚未建立 Kong GuZheng 状态文件",
        ],
    },
}


def parse_args():
    repo_root = Path(__file__).resolve().parents[1]
    workspace_root = repo_root.parent
    default_docx = workspace_root / "doc" / DOC_NAME
    default_assets = workspace_root / "mgsc_daw_assets"
    default_output = repo_root / "config" / OUTPUT_NAME

    parser = argparse.ArgumentParser(
        description="从云端 DAW 需求文档生成 MIDI 音源映射 JSON 配置。"
    )
    parser.add_argument("--docx", default=str(default_docx), help="需求文档 docx 路径")
    parser.add_argument("--assets", default=str(default_assets), help="本地音源资产目录")
    parser.add_argument("--output", default=str(default_output), help="输出 JSON 路径")
    parser.add_argument(
        "--table-index",
        type=int,
        default=4,
        help="Word 表格索引，默认读取第 4 个大表，索引从 0 开始",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="只打印摘要，不写入 JSON 文件",
    )
    return parser.parse_args()


def load_document(path):
    try:
        from docx import Document
    except ImportError:
        print("缺少 python-docx，请先安装：python -m pip install python-docx", file=sys.stderr)
        raise
    return Document(str(path))


def cell_text(cell):
    return cell.text.strip().replace("\n", " / ")


def int_or_none(value):
    value = str(value).strip()
    if not value:
        return None
    try:
        return int(value)
    except ValueError:
        return None


def normalize_slashes(value):
    return " ".join(str(value).replace("/", " ").split())


def preset_relpath(plugin_name, bank, program):
    if plugin_name == "Musyng_Kite":
        return "soundfont2/Musyng_Kite.sf2"
    if plugin_name == "Keyzone Classic":
        return "Steinberg/VstPlugins/Keyzone Classic/Keyzone Classic/{0}.txt".format(program)
    if plugin_name == "DSK Saxophones":
        return "Steinberg/VstPlugins/DSK Saxophones/DSK Saxophones/{0}.txt".format(program)
    if plugin_name == "Sonatina Orchestra":
        return "Steinberg/VstPlugins/Sonatina Orchestra/Sonatina Orchestra/{0}/{1}.txt".format(
            bank,
            program,
        )
    return None


def build_mapping(row_index, cells, assets_root):
    category = cells[0]
    user_name = cells[1]
    midi_id = cells[2]
    midi_name = cells[3]
    web_bank = cells[4]
    web_program = cells[5]
    plugin_name = cells[6]
    plugin_type = cells[7]
    cloud_bank = cells[8]
    cloud_program = cells[9]

    plugin_info = PLUGIN_INFO.get(plugin_name, {})
    override = COMPATIBILITY_OVERRIDES.get((midi_id, plugin_name), {})

    normalized_bank = override.get("normalized_bank", cloud_bank)
    normalized_program = override.get("normalized_program", cloud_program)
    if "normalized_program" not in override and "/" in normalized_program:
        normalized_program = normalize_slashes(normalized_program)

    target_preset = override.get(
        "local_preset_relpath",
        preset_relpath(plugin_name, normalized_bank, normalized_program),
    )
    candidate_relpaths = override.get("candidate_relpaths")
    if not candidate_relpaths and target_preset:
        candidate_relpaths = [target_preset]

    candidate_status = []
    for relpath in candidate_relpaths or []:
        local_path = assets_root / Path(relpath.replace("/", "\\"))
        candidate_status.append(
            {
                "relpath": relpath,
                "local_path": str(local_path),
                "exists": local_path.exists(),
            }
        )

    needs_confirmation = list(override.get("needs_confirmation", []))
    notes = list(override.get("notes", []))
    if cloud_bank == "？":
        needs_confirmation.append("文档 cloud bank 为 ？，运行配置中需要表达为插件预设而不是 SF2 Bank")

    is_drum = category == "节奏轨鼓组"
    mapping_id = "drum_128_{0:03d}".format(int_or_none(midi_id) or 0) if is_drum else "gm_{0:03d}".format(int_or_none(midi_id) or 0)

    return {
        "id": mapping_id,
        "source_doc": {
            "table_row": row_index,
            "category": category,
            "user_instrument_name": user_name,
            "midi_id": int_or_none(midi_id),
            "midi_name": midi_name,
            "web_bank": int_or_none(web_bank),
            "web_program": int_or_none(web_program),
            "cloud_plugin_name": plugin_name,
            "cloud_type_raw": plugin_type,
            "cloud_bank_raw": cloud_bank,
            "cloud_program_raw": cloud_program,
        },
        "midi": {
            "bank": int_or_none(web_bank),
            "program": int_or_none(web_program),
            "is_drum_bank": is_drum,
        },
        "target": {
            "plugin_name": plugin_name,
            "plugin_id": plugin_info.get("plugin_id", plugin_name),
            "plugin_type": plugin_info.get("plugin_type", plugin_type.lower()),
            "plugin_asset_relpath": plugin_info.get("asset_relpath"),
            "bank": normalized_bank,
            "program": normalized_program,
            "preset_candidates": candidate_status,
        },
        "status": {
            "implementation": plugin_info.get("implementation_status", "planned"),
            "needs_confirmation": needs_confirmation,
            "notes": notes,
        },
    }


def build_config(docx_path, assets_root, table_index):
    document = load_document(docx_path)
    if table_index >= len(document.tables):
        raise RuntimeError(
            "文档只有 {0} 个表格，无法读取 table-index={1}".format(
                len(document.tables),
                table_index,
            )
        )

    table = document.tables[table_index]
    mappings = []
    for row_index, row in enumerate(table.rows[2:], start=2):
        cells = [cell_text(cell) for cell in row.cells]
        if not any(cells):
            continue
        mappings.append(build_mapping(row_index, cells, assets_root))

    plugin_counts = Counter(item["target"]["plugin_name"] for item in mappings)
    confirmation_count = sum(
        1 for item in mappings if item["status"]["needs_confirmation"]
    )

    return {
        "_file_info": {
            "File name": OUTPUT_NAME,
            "Brief": "云端 DAW MIDI 音源映射配置",
            "Function": "记录需求文档中 MIDI Bank、Program 到云端插件和预设的结构化映射",
            "Author": "咪咕数创工程架构组 / MGSC AI Software Architecture group",
            "Version": "V2.5.10",
            "Date": date.today().strftime("%Y/%m/%d"),
        },
        "source": {
            "docx": str(docx_path),
            "table_index": table_index,
            "note": "最终需求以云端DAW音频工作站引擎.docx为准，本配置保留兼容修正和待确认项。",
        },
        "summary": {
            "mapping_count": len(mappings),
            "normal_gm_count": sum(
                1 for item in mappings if not item["midi"]["is_drum_bank"]
            ),
            "drum_bank_count": sum(
                1 for item in mappings if item["midi"]["is_drum_bank"]
            ),
            "plugin_counts": dict(sorted(plugin_counts.items())),
            "needs_confirmation_count": confirmation_count,
        },
        "plugins": PLUGIN_INFO,
        "mappings": mappings,
    }


def print_summary(config):
    summary = config["summary"]
    print("mapping_count:", summary["mapping_count"])
    print("normal_gm_count:", summary["normal_gm_count"])
    print("drum_bank_count:", summary["drum_bank_count"])
    print("needs_confirmation_count:", summary["needs_confirmation_count"])
    print("plugin_counts:")
    for name, count in summary["plugin_counts"].items():
        print("  {0}: {1}".format(name, count))


def main():
    args = parse_args()
    docx_path = Path(args.docx)
    assets_root = Path(args.assets)
    output_path = Path(args.output)

    config = build_config(docx_path, assets_root, args.table_index)
    print_summary(config)

    if args.dry_run:
        return 0

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(
        json.dumps(config, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    print("wrote:", str(output_path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
