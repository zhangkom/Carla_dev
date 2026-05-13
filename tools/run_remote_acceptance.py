#!/usr/bin/env python3
# /**
# * File name: run_remote_acceptance.py
# * Brief: 远程 Ubuntu 服务端渲染验收脚本
# * Function:
# *     从 Windows 客户端批量调用远程 /mgsc_daw_service/v1/render，
# *     保存响应、MP3、音量/耗时统计，并生成 Markdown/JSON/DOCX 报告。
# * Author: 软件工程架构组
# *     MGSC AI Software Architecture group
# * Version: V2.5.10
# * Date: 2026/05/11
# */
from __future__ import annotations

import argparse
import base64
import json
import math
import mimetypes
import re
import shutil
import subprocess
import sys
import time
import uuid
import zipfile
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
from typing import Any
from urllib import request, error


@dataclass
class ZipCase:
    name: str
    path: str
    size_bytes: int
    midi_count: int
    conf_count: int
    nested_zip_count: int
    valid_render_bundle: bool
    skip_reason: str | None = None


@dataclass
class CaseResult:
    zip_name: str
    ok: bool
    skipped: bool
    seconds: float | None = None
    status_code: int | None = None
    error: str | None = None
    job_id: str | None = None
    style_id: str | None = None
    mode: str | None = None
    route_count: int | None = None
    mp3_path: str | None = None
    response_path: str | None = None
    stats_path: str | None = None
    mp3_bytes: int | None = None
    duration_seconds: float | None = None
    mean_volume_db: float | None = None
    max_volume_db: float | None = None
    renderer_total_seconds: float | None = None
    record_audio_seconds: float | None = None
    ffmpeg_mp3_seconds: float | None = None
    request_total_seconds: float | None = None
    artifact_archive: Any | None = None
    timing_summary: Any | None = None


def _log(log_path: Path, message: str) -> None:
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    line = f"{timestamp} {message}"
    print(line, flush=True)
    with log_path.open("a", encoding="utf-8") as handle:
        handle.write(line + "\n")


def inspect_zip(path: Path) -> ZipCase:
    with zipfile.ZipFile(path) as bundle:
        names = bundle.namelist()
    midi_count = sum(name.lower().endswith((".mid", ".midi")) for name in names)
    conf_count = sum(name.lower().endswith("conf.json") for name in names)
    nested_zip_count = sum(name.lower().endswith(".zip") for name in names)
    valid = midi_count > 0 and conf_count > 0
    reason = None if valid else "not a render bundle: missing MIDI or conf.json"
    return ZipCase(
        name=path.name,
        path=str(path),
        size_bytes=path.stat().st_size,
        midi_count=midi_count,
        conf_count=conf_count,
        nested_zip_count=nested_zip_count,
        valid_render_bundle=valid,
        skip_reason=reason,
    )


def multipart_upload(url: str, field_name: str, file_path: Path, timeout: float) -> tuple[int, bytes]:
    boundary = "----mgscdaw" + uuid.uuid4().hex
    content_type = mimetypes.guess_type(file_path.name)[0] or "application/octet-stream"
    head = (
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="{field_name}"; filename="{file_path.name}"\r\n'
        f"Content-Type: {content_type}\r\n\r\n"
    ).encode("utf-8")
    tail = f"\r\n--{boundary}--\r\n".encode("utf-8")
    body = head + file_path.read_bytes() + tail
    req = request.Request(
        url,
        data=body,
        method="POST",
        headers={
            "Content-Type": f"multipart/form-data; boundary={boundary}",
            "Content-Length": str(len(body)),
        },
    )
    try:
        with request.urlopen(req, timeout=timeout) as response:
            return response.status, response.read()
    except error.HTTPError as exc:
        return exc.code, exc.read()


def safe_stem(name: str) -> str:
    text = Path(name).stem
    text = re.sub(r"[^\w.-]+", "_", text, flags=re.UNICODE).strip("._")
    return text or "render"


def decode_mp3(payload: dict[str, Any], output_path: Path) -> int:
    mp3_file = payload.get("mp3_file")
    if not isinstance(mp3_file, dict):
        raise RuntimeError("response does not include mp3_file")
    encoded = mp3_file.get("base64")
    if not isinstance(encoded, str) or not encoded:
        raise RuntimeError("response does not include mp3_file.base64")
    raw = base64.b64decode(encoded.encode("ascii"), validate=True)
    output_path.write_bytes(raw)
    return len(raw)


def parse_float(value: str) -> float | None:
    if value == "-inf":
        return -math.inf
    try:
        return float(value)
    except ValueError:
        return None


def audio_stats(mp3_path: Path, ffmpeg: str, ffprobe: str) -> dict[str, Any]:
    stats: dict[str, Any] = {}
    duration = subprocess.run(
        [
            ffprobe,
            "-v",
            "error",
            "-show_entries",
            "format=duration",
            "-of",
            "default=noprint_wrappers=1:nokey=1",
            str(mp3_path),
        ],
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        check=False,
    )
    if duration.returncode == 0:
        try:
            stats["duration_seconds"] = round(float(duration.stdout.strip()), 3)
        except ValueError:
            stats["duration_seconds"] = None
    else:
        stats["duration_error"] = duration.stderr.strip()

    volume = subprocess.run(
        [
            ffmpeg,
            "-hide_banner",
            "-nostats",
            "-i",
            str(mp3_path),
            "-af",
            "volumedetect",
            "-f",
            "null",
            "-",
        ],
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        check=False,
    )
    text = volume.stdout + "\n" + volume.stderr
    for key in ("mean_volume", "max_volume"):
        match = re.search(rf"{key}:\s+(-?inf|-?\d+(?:\.\d+)?) dB", text)
        if match:
            stats[key + "_db"] = parse_float(match.group(1))
    if volume.returncode != 0:
        stats["volume_error"] = text.strip()
    stats["non_silent"] = (
        isinstance(stats.get("max_volume_db"), (float, int))
        and not math.isinf(float(stats["max_volume_db"]))
        and float(stats["max_volume_db"]) > -80.0
    )
    return stats


def extract_timing(payload: dict[str, Any], result: CaseResult) -> None:
    summary = payload.get("timing_summary")
    result.timing_summary = summary
    if isinstance(summary, dict):
        result.renderer_total_seconds = summary.get("renderer_total_seconds")
        result.record_audio_seconds = summary.get("record_audio_seconds")
        result.ffmpeg_mp3_seconds = summary.get("ffmpeg_mp3_seconds")
        result.request_total_seconds = summary.get("request_total_seconds")
    timings = payload.get("timings")
    if isinstance(timings, dict):
        result.request_total_seconds = result.request_total_seconds or timings.get("request_total_seconds")
    renderer_stage = payload.get("renderer_stage_seconds")
    if isinstance(renderer_stage, dict):
        result.record_audio_seconds = result.record_audio_seconds or renderer_stage.get("record_audio_seconds")
        result.ffmpeg_mp3_seconds = result.ffmpeg_mp3_seconds or renderer_stage.get("ffmpeg_mp3_seconds")


def infer_mode_and_route_count(payload: dict[str, Any]) -> tuple[str | None, int | None]:
    auto_route = payload.get("auto_route")
    if isinstance(auto_route, dict):
        return auto_route.get("mode") or "auto", auto_route.get("route_count") or 1
    midi_policy = payload.get("midi_policy")
    if isinstance(midi_policy, dict):
        return midi_policy.get("mode"), midi_policy.get("route_count")
    return None, None


def run_case(
    case: ZipCase,
    args: argparse.Namespace,
    output_dir: Path,
    log_path: Path,
) -> CaseResult:
    response_dir = output_dir / "responses"
    mp3_dir = output_dir / "mp3"
    stats_dir = output_dir / "stats"
    for directory in (response_dir, mp3_dir, stats_dir):
        directory.mkdir(parents=True, exist_ok=True)

    if not case.valid_render_bundle:
        _log(log_path, f"SKIP {case.name}: {case.skip_reason}")
        return CaseResult(zip_name=case.name, ok=False, skipped=True, error=case.skip_reason)

    zip_path = Path(case.path)
    _log(log_path, f"START {case.name}")
    started = time.monotonic()
    result = CaseResult(zip_name=case.name, ok=False, skipped=False)
    response_path = response_dir / f"{safe_stem(case.name)}.json"
    result.response_path = str(response_path)
    try:
        status_code, raw = multipart_upload(args.url, args.field, zip_path, args.timeout)
        result.status_code = status_code
        result.seconds = round(time.monotonic() - started, 3)
        response_path.write_bytes(raw)
        if status_code < 200 or status_code >= 300:
            text = raw[:2000].decode("utf-8", errors="replace")
            raise RuntimeError(f"HTTP {status_code}: {text}")
        payload = json.loads(raw.decode("utf-8"))
        result.job_id = payload.get("job_id")
        result.style_id = payload.get("style_id")
        result.artifact_archive = payload.get("artifact_archive")
        result.mode, result.route_count = infer_mode_and_route_count(payload)
        extract_timing(payload, result)

        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        mp3_path = mp3_dir / f"{safe_stem(case.name)}_{args.version}_{stamp}.mp3"
        result.mp3_bytes = decode_mp3(payload, mp3_path)
        result.mp3_path = str(mp3_path)
        stats = audio_stats(mp3_path, args.ffmpeg, args.ffprobe)
        stats_path = stats_dir / f"{safe_stem(case.name)}_audio_stats.json"
        stats_path.write_text(json.dumps(stats, ensure_ascii=False, indent=2), encoding="utf-8")
        result.stats_path = str(stats_path)
        result.duration_seconds = stats.get("duration_seconds")
        result.mean_volume_db = stats.get("mean_volume_db")
        result.max_volume_db = stats.get("max_volume_db")
        result.ok = bool(stats.get("non_silent"))
        state = "PASS" if result.ok else "FAIL"
        _log(
            log_path,
            (
                f"{state} {case.name}: seconds={result.seconds} "
                f"style={result.style_id} route_count={result.route_count} "
                f"mean={result.mean_volume_db} max={result.max_volume_db}"
            ),
        )
    except Exception as exc:
        result.seconds = result.seconds or round(time.monotonic() - started, 3)
        result.error = str(exc)
        _log(log_path, f"FAIL {case.name}: {result.error}")
    return result


def markdown_table(results: list[CaseResult], output_dir: Path) -> str:
    lines = [
        "| zip | result | seconds | duration | mean dB | max dB | style | routes | mp3 | error |",
        "| --- | --- | ---: | ---: | ---: | ---: | --- | ---: | --- | --- |",
    ]
    for item in results:
        if item.skipped:
            state = "SKIP"
        else:
            state = "PASS" if item.ok else "FAIL"
        mp3 = ""
        if item.mp3_path:
            try:
                mp3 = str(Path(item.mp3_path).relative_to(output_dir))
            except ValueError:
                mp3 = item.mp3_path
        lines.append(
            "| {zip} | {state} | {seconds} | {duration} | {mean} | {maxv} | {style} | {routes} | {mp3} | {error} |".format(
                zip=item.zip_name,
                state=state,
                seconds="" if item.seconds is None else item.seconds,
                duration="" if item.duration_seconds is None else item.duration_seconds,
                mean="" if item.mean_volume_db is None else item.mean_volume_db,
                maxv="" if item.max_volume_db is None else item.max_volume_db,
                style=item.style_id or "",
                routes="" if item.route_count is None else item.route_count,
                mp3=mp3,
                error=(item.error or "").replace("|", "/"),
            )
        )
    return "\n".join(lines)


def write_reports(
    args: argparse.Namespace,
    output_dir: Path,
    cases: list[ZipCase],
    results: list[CaseResult],
) -> Path:
    passed = sum(1 for item in results if item.ok)
    failed = sum(1 for item in results if not item.ok and not item.skipped)
    skipped = sum(1 for item in results if item.skipped)
    summary = {
        "version": args.version,
        "render_url": args.url,
        "zip_dir": str(Path(args.zip_dir).resolve()),
        "output_dir": str(output_dir),
        "started_at": args.started_at,
        "finished_at": datetime.now().isoformat(timespec="seconds"),
        "passed": passed,
        "failed": failed,
        "skipped": skipped,
        "cases": [asdict(case) for case in cases],
        "results": [asdict(result) for result in results],
    }
    (output_dir / "summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")

    md = "\n".join(
        [
            f"# {args.version} 远程接口验收测试",
            "",
            f"- 测试时间：{summary['started_at']} - {summary['finished_at']}",
            f"- 客户端：Windows",
            f"- 服务端接口：`{args.url}`",
            f"- ZIP 目录：`{summary['zip_dir']}`",
            f"- 输出目录：`{summary['output_dir']}`",
            f"- 结果：通过 {passed}，失败 {failed}，跳过 {skipped}",
            "",
            "## 测试结果",
            "",
            markdown_table(results, output_dir),
            "",
            "## 判定规则",
            "",
            "- HTTP 2xx 且响应包含 `mp3_file.base64`。",
            "- MP3 可解码，`ffmpeg volumedetect` 的 `max_volume` 大于 -80 dB 判定为非静音。",
            "- `daojianrumeng_0508_test_zips.zip` 这类内部只包含 zip 的集合包不作为单个渲染输入。",
            "",
            "## 当前结论",
            "",
            "- `vst_keyzone_classic` 在性能分支上通过 `MUSIC_SERVICE_DUMMY_SLEEP_DIVISOR_BY_PLUGIN=vst_keyzone_classic=16` 和 2 秒预热走加速路径；如清空该变量，仍可由 `MUSIC_SERVICE_DUMMY_NOSLEEP_DISABLE_PLUGINS=vst_keyzone_classic` 回退实时模式。",
            "- Kong Audio 仍保持 nosleep 加速路径，当前远程测试应重点观察是否仍在几十秒级。",
        ]
    )
    summary_md = output_dir / "SUMMARY.md"
    summary_md.write_text(md, encoding="utf-8")

    report_docx = Path(args.doc_dir) / f"远程接口验收测试报告_{args.version}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    try:
        from docx import Document
        from docx.shared import Inches

        doc = Document()
        doc.add_heading("远程接口验收测试报告", level=1)
        doc.add_paragraph(f"版本：{args.version}")
        doc.add_paragraph(f"测试时间：{summary['started_at']} - {summary['finished_at']}")
        doc.add_paragraph(f"客户端：Windows")
        doc.add_paragraph(f"服务端接口：{args.url}")
        doc.add_paragraph(f"ZIP 目录：{summary['zip_dir']}")
        doc.add_paragraph(f"输出目录：{summary['output_dir']}")
        doc.add_paragraph(f"汇总：通过 {passed}，失败 {failed}，跳过 {skipped}")
        doc.add_heading("测试结果", level=2)
        table = doc.add_table(rows=1, cols=9)
        table.style = "Table Grid"
        headers = ["zip", "结果", "耗时(s)", "音频时长(s)", "mean dB", "max dB", "style", "routes", "备注"]
        for cell, header in zip(table.rows[0].cells, headers):
            cell.text = header
        for item in results:
            row = table.add_row().cells
            row[0].text = item.zip_name
            row[1].text = "SKIP" if item.skipped else ("PASS" if item.ok else "FAIL")
            row[2].text = "" if item.seconds is None else str(item.seconds)
            row[3].text = "" if item.duration_seconds is None else str(item.duration_seconds)
            row[4].text = "" if item.mean_volume_db is None else str(item.mean_volume_db)
            row[5].text = "" if item.max_volume_db is None else str(item.max_volume_db)
            row[6].text = item.style_id or ""
            row[7].text = "" if item.route_count is None else str(item.route_count)
            row[8].text = item.error or ""
        doc.add_heading("判定规则", level=2)
        doc.add_paragraph("HTTP 2xx 且响应包含 mp3_file.base64；MP3 可解码；max_volume 大于 -80 dB 判定为非静音。")
        doc.add_heading("建议", level=2)
        doc.add_paragraph("继续基于需求文档构造 137 条 Bank/Program 覆盖包，先按云端目标音源和边界分层，再按每条映射生成完整测试包。")
        doc.save(report_docx)
    except Exception as exc:
        (output_dir / "DOCX_REPORT_ERROR.txt").write_text(str(exc), encoding="utf-8")
    return summary_md


def main() -> int:
    parser = argparse.ArgumentParser(description="Run remote MGSC DAW render acceptance tests.")
    parser.add_argument("--url", required=True)
    parser.add_argument("--zip-dir", required=True)
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--doc-dir", required=True)
    parser.add_argument("--field", default="data")
    parser.add_argument("--version", default="6.5.12.1508")
    parser.add_argument("--timeout", type=float, default=1200.0)
    parser.add_argument("--ffmpeg", default=shutil.which("ffmpeg") or "ffmpeg")
    parser.add_argument("--ffprobe", default=shutil.which("ffprobe") or "ffprobe")
    args = parser.parse_args()
    args.started_at = datetime.now().isoformat(timespec="seconds")

    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    Path(args.doc_dir).mkdir(parents=True, exist_ok=True)
    log_path = output_dir / "run.log"

    zip_dir = Path(args.zip_dir)
    cases = [inspect_zip(path) for path in sorted(zip_dir.glob("*.zip"))]
    _log(log_path, f"Remote acceptance started url={args.url}")
    _log(log_path, f"Found {len(cases)} zip files")
    results = []
    for case in cases:
        result = run_case(case, args, output_dir, log_path)
        results.append(result)
    summary_md = write_reports(args, output_dir, cases, results)
    _log(log_path, f"Remote acceptance finished summary={summary_md}")
    failed = sum(1 for item in results if not item.ok and not item.skipped)
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
